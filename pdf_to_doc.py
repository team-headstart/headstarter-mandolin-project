import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import tempfile
from dotenv import load_dotenv
from PIL import Image

# Load environment variables
load_dotenv()

# Try to import Google AI packages
try:
    import google.generativeai as genai
    from google.generativeai.types import HarmCategory, HarmBlockThreshold
    GOOGLE_AI_AVAILABLE = True
except ImportError:
    print("Warning: Google Generative AI package not found. Falling back to Tesseract OCR.")
    GOOGLE_AI_AVAILABLE = False

def setup_gemini():
    """Setup Gemini AI client with safety settings"""
    if not GOOGLE_AI_AVAILABLE:
        return None
        
    try:
        api_key = os.getenv('GOOGLE_API_KEY')
        if not api_key:
            print("Warning: GOOGLE_API_KEY not found in environment variables")
            return None
            
        genai.configure(api_key=api_key)
        
        # Configure safety settings
        safety_settings = {
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
        }
        
        return genai, safety_settings
    except Exception as e:
        print(f"Error setting up Gemini AI: {str(e)}")
        return None

def process_with_gemini(client, image_path, prompt="Extract and validate the text from this image, focusing on phone numbers, fax numbers, and comments"):
    """Process image using Gemini AI with enhanced error handling"""
    if not client:
        return None
        
    try:
        genai_client, safety_settings = client
        
        # Load the image using PIL
        image = Image.open(image_path)
        
        # Generate content using Gemini with safety settings
        model = genai_client.GenerativeModel('gemini-1.5-flash', safety_settings=safety_settings)
        
        # Configure generation parameters
        generation_config = {
            "temperature": 0.4,
            "top_p": 0.8,
            "top_k": 40,
            "max_output_tokens": 2048,
        }
        
        response = model.generate_content(
            [prompt, image],
            generation_config=generation_config
        )
        
        if response.prompt_feedback.block_reason:
            print(f"Warning: Content blocked due to: {response.prompt_feedback.block_reason}")
            return None
            
        return response.text
    except Exception as e:
        print(f"Error processing with Gemini: {str(e)}")
        return None

def validate_phone_number(text):
    """Validate and format phone number"""
    # Remove any non-digit characters
    digits = re.sub(r'\D', '', text)
    # Check if it's a valid phone number (7-15 digits)
    if 7 <= len(digits) <= 15:
        return True
    return False

def validate_fax_number(text):
    """Validate and format fax number"""
    # Similar to phone number validation
    digits = re.sub(r'\D', '', text)
    if 7 <= len(digits) <= 15:
        return True
    return False

def format_phone_number(text):
    """Format phone number with proper spacing"""
    digits = re.sub(r'\D', '', text)
    if len(digits) == 10:
        return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    return text

def process_text(text):
    """Process and validate text content"""
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check for phone/fax numbers
        if 'phone' in line.lower() or 'tel' in line.lower():
            parts = line.split(':')
            if len(parts) > 1:
                label, number = parts[0], parts[1].strip()
                if validate_phone_number(number):
                    processed_lines.append(f"{label}: {format_phone_number(number)}")
                else:
                    processed_lines.append(f"{label}: [Invalid phone number] {number}")
            else:
                processed_lines.append(line)
                
        elif 'fax' in line.lower():
            parts = line.split(':')
            if len(parts) > 1:
                label, number = parts[0], parts[1].strip()
                if validate_fax_number(number):
                    processed_lines.append(f"{label}: {format_phone_number(number)}")
                else:
                    processed_lines.append(f"{label}: [Invalid fax number] {number}")
            else:
                processed_lines.append(line)
                
        # Handle comments section
        elif 'comment' in line.lower():
            processed_lines.append("\nComments:")
            processed_lines.append("-" * 40)
            
        else:
            processed_lines.append(line)
    
    return '\n'.join(processed_lines)

def convert_pdf_to_doc(pdf_path=None, output_doc_path=None):
    """
    Convert handwritten PDF to Word document with validation using Gemini AI
    
    Args:
        pdf_path (str): Path to the input PDF file
        output_doc_path (str): Path where the output Word document will be saved
    """
    try:
        # Use environment variables if paths not provided
        pdf_path = pdf_path or os.getenv('DEFAULT_INPUT_PDF', 'input.pdf')
        output_doc_path = output_doc_path or os.getenv('DEFAULT_OUTPUT_DOC', 'output.docx')
        
        # Validate input file exists
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Input PDF file not found: {pdf_path}")
        
        # Setup Gemini AI client
        use_gemini = os.getenv('USE_GEMINI', 'true').lower() == 'true' and GOOGLE_AI_AVAILABLE
        fallback_to_tesseract = os.getenv('FALLBACK_TO_TESSERACT', 'true').lower() == 'true'
        
        client = None
        if use_gemini:
            client = setup_gemini()
            if not client and fallback_to_tesseract:
                print("Failed to setup Gemini AI. Falling back to Tesseract OCR.")
            elif not client:
                raise Exception("Failed to setup Gemini AI and fallback is disabled.")
        
        # Create a new Word document
        doc = Document()
        
        # Set default font from environment variables
        style = doc.styles['Normal']
        style.font.name = os.getenv('DEFAULT_FONT', 'Arial')
        style.font.size = Pt(int(os.getenv('DEFAULT_FONT_SIZE', '11')))
        
        # Convert PDF to images
        print("Converting PDF to images...")
        try:
            images = convert_from_path(pdf_path)
        except Exception as e:
            raise Exception(f"Failed to convert PDF to images: {str(e)}")
        
        # Process each page
        for i, image in enumerate(images):
            print(f"Processing page {i+1}...")
            
            # Add page number
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header.add_run(f"Page {i+1}")
            
            # Save image temporarily
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                image.save(temp_file.name, 'PNG')
                temp_path = temp_file.name
            
            try:
                if client:
                    # Process with Gemini AI
                    text = process_with_gemini(client, temp_path)
                    if not text and fallback_to_tesseract:
                        print(f"Gemini AI processing failed for page {i+1}. Falling back to Tesseract OCR.")
                        text = pytesseract.image_to_string(image)
                    elif not text:
                        raise Exception(f"Failed to process page {i+1} with Gemini AI and fallback is disabled.")
                else:
                    # Use Tesseract OCR
                    text = pytesseract.image_to_string(image)
                
                if not text.strip():
                    print(f"Warning: No text extracted from page {i+1}")
                    continue
                
                # Process and validate the text
                processed_text = process_text(text)
                
                # Add text to document
                paragraph = doc.add_paragraph()
                paragraph.add_run(processed_text)
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    print(f"Warning: Failed to delete temporary file: {str(e)}")
            
            # Add page break if not the last page
            if i < len(images) - 1:
                doc.add_page_break()
        
        # Save the document
        try:
            doc.save(output_doc_path)
            print(f"Document saved successfully at: {output_doc_path}")
        except Exception as e:
            raise Exception(f"Failed to save document: {str(e)}")
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        raise

if __name__ == "__main__":
    try:
        # Get paths from environment variables or use defaults
        pdf_path = os.getenv('DEFAULT_INPUT_PDF', 'input.pdf')
        output_doc_path = os.getenv('DEFAULT_OUTPUT_DOC', 'output.docx')
        
        convert_pdf_to_doc(pdf_path, output_doc_path)
    except Exception as e:
        print(f"Fatal error: {str(e)}")
        exit(1) 